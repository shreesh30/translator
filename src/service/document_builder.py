import re
from typing import List, Dict
import logging
from PIL import ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.model.footer import Footer
from src.model.language_config import LanguageConfig
from src.model.paragraph import Paragraph
from src.service.document_processor import DocumentProcessor
from src.utils.utils import Utils


class DocumentBuilder:
    PAGE_USED = 0

    def __init__(self, language_config: LanguageConfig, document_processor: DocumentProcessor, document: Document):
        self.logger = logging.getLogger(__name__)

        self.document_processor = document_processor
        self.language_config = language_config
        self.document = document
        self.font_name = self.language_config.get_target_font_name()
        self._configure_docx_section(self.document.sections[-1])

    def build_document(self, paragraphs: List[Paragraph]):
        self.add_paragraphs(paragraphs)
        self.add_page_numbers()

    def add_paragraphs(self, paragraphs: List[Paragraph]):
        pages = self.document_processor.get_pages()
        _, page_number_start = self.document_processor.get_page_number_info()
        header_page_number_start = None

        if page_number_start is not None:
            header_page_number_start = page_number_start + 1

        for paragraph in paragraphs:
            self.logger.info(f'Adding Paragraph:{paragraph}')

            page = pages[paragraph.get_page_number()]

            lines = [line.get_text() for line in paragraph.get_lines()]
            current_para = None
            continue_para = False

            has_footer = bool(paragraph.get_footer())
            # Flag to track if footer for current paragraph has been added to a section
            footer_added = False

            for i, line in enumerate(lines):
                # Calculate height of current line
                line_spacing_in = (self.language_config.get_line_spacing_multiplier() * self.language_config.get_font_size_multiplier() * paragraph.get_font_size())/72
                line_height = self.estimate_line_height(line, paragraph)
                line_height_in = line_height/72

                predicted_space_used = self.PAGE_USED+line_spacing_in+line_height_in

                # Check if we need to create a new section (page) or paragraph
                if predicted_space_used > Utils.USABLE_PAGE_HEIGHT:
                    self.logger.info('Creating New Section')
                    # If this is the last para in the section, remove space after for it before creating a new section
                    if self.document.paragraphs:
                        current_para_format = self.document.paragraphs[-1].paragraph_format
                        current_para_format.space_after = Pt(0)

                    # If footer is pending for current paragraph, attach to current_section BEFORE making a new one
                    if has_footer and not footer_added:
                        self._add_footer(self.document.sections[-1], paragraph.get_footer())
                        footer_added = True

                    # Create new section (page)
                    self.document.add_section(WD_SECTION.NEW_PAGE)
                    self._configure_docx_section(self.document.sections[-1])
                    self.PAGE_USED = 0
                    current_para = None  # Force new paragraph on new page

                    if i!=0: #If it is not the first line of the paragraph on the new page, set the flag
                        continue_para = True

                # Create new paragraph if needed (first line or line doesn't fit in current para)
                if current_para is None:
                    self.logger.info('Creating New Paragraph')
                    current_para = self.document.add_paragraph()
                    self.set_rtl(current_para)
                    self._set_paragraph_alignment_and_indent(current_para, paragraph, page, self.document.sections[-1], continue_para)
                    self._set_paragraph_spacing(current_para, paragraph)

                    if has_footer and not footer_added:
                        self._add_footer(self.document.sections[-1], paragraph.get_footer())
                        footer_added = True

                # Add the line to current paragraph
                self._add_text_with_styling(current_para, line, paragraph)
                self.PAGE_USED += line_height_in

            if has_footer and not footer_added:
                self._add_footer(self.document.sections[-1], paragraph.get_footer())

            if header_page_number_start is not None and len(self.document.sections) - 1 >= header_page_number_start:
                section_index =  len(self.document.sections) - 1
                is_volume = (section_index - header_page_number_start) % 2 == 0
                self.add_header(paragraph, self.document.sections[-1], is_volume)

            # TODO: CHECK HOW CAN WE AVOID ADDING SPACE AFTER FOR LAST LINE IN THE SECTION(CHECK IF THERE IS SPACE FOR ADDING 1 LINE OR NOT, IF THERE IS NOT THEN SKIP ADDING SPACE_AFTER ELSE ADD SPACE_AFTER)
            self.PAGE_USED += ((paragraph.get_font_size() * 0.5) / 72)  # Adding to incorporate "Space After" after a paragraph

            # Process sub-paragraphs
            if paragraph.get_sub_paragraphs():
                self.add_paragraphs(paragraph.get_sub_paragraphs())

    @staticmethod
    def clear_header(header):
        for para in header.paragraphs:
            p = para._element
            p.getparent().remove(p)

    def get_text_width(self,text, font_size_pt):
        font = ImageFont.truetype(self.language_config.get_target_font_path(), size=int(font_size_pt))
        return font.getlength(text)

    def add_header(self, paragraph, section, is_volume):
        self.logger.info(f'Processing Header: Paragraph: {paragraph}, Is Volume: {is_volume}')
        header = section.first_page_header
        header.is_linked_to_previous = False
        self.clear_header(header)

        header_para = header.add_paragraph()

        # Get header text based on type
        header_text = paragraph.get_volume() if is_volume else paragraph.get_chapter()
        self.logger.info(f'Header Text: {header_text}')

        if not header_text or not header_text.strip():
            return

        if is_volume:
            # Reformat header text to insert tab stops
            parts = header_text.split("[", 1)
            if len(parts) == 2:
                left, right = parts[0].strip(), "[" + parts[1]
                usable_width = Utils.STANDARDIZED_PAGE_WIDTH - Utils.STANDARDIZED_LEFT_MARGIN - Utils.STANDARDIZED_RIGHT_MARGIN
                center_tab = usable_width / 2
                right_tab = usable_width-center_tab- self.get_text_width(left, 9 * self.language_config.get_font_size_multiplier())
                header_text = f"\t{left}\t{right}"

                header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                tab_stops = header_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(center_tab), alignment=WD_TAB_ALIGNMENT.CENTER)
                tab_stops.add_tab_stop(Inches(right_tab/2), alignment=WD_TAB_ALIGNMENT.RIGHT)
        else:
            # Center-align chapter text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add text as a run
        self.logger.info(f'Adding Header: {header_text}')
        run = header_para.add_run(header_text)
        run.font.name = self.font_name
        run.font.size = Pt(9 * self.language_config.get_font_size_multiplier())
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)

    def _add_footer(self, section, para_footers: List[Footer]):
        """Adds footer ONLY to the current section's next page"""
        self.logger.info(f'Adding Footers: {para_footers}')
        # 1. Get current section (always use last section)
        section.footer_distance = Inches(Utils.STANDARDIZED_FOOTER_DISTANCE)
        footer = section.first_page_footer
        footer.is_linked_to_previous = False

        for para_footer in para_footers:
            # 2. Add horizontal line only if this is the FIRST footer paragraph
            if not any(p.text.strip() for p in footer.paragraphs):
                line_para = footer.add_paragraph()
                self.add_horizontal_line(line_para)
                line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                line_para_format = line_para.paragraph_format
                line_para_format.left_indent = Pt(0)
                line_para_format.right_indent = Pt(0)
                line_para_format.space_before = Pt(0)
                line_para_format.space_after = Pt(0)
                line_para_format.line_spacing = Pt(0)

            if para_footer.get_text().strip():  # Only add if there's actual text
                footer_para = footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Or whatever alignment is desired
                footer_format = footer_para.paragraph_format
                footer_format.space_before = Pt(0)
                footer_format.space_after = Pt(0)
                footer_format.line_spacing = Pt(
                    para_footer.get_font_size() * self.language_config.get_font_size_multiplier() * self.language_config.get_line_spacing_multiplier())
                run = footer_para.add_run(para_footer.get_text())
                run.font.size = Pt(para_footer.get_font_size() * self.language_config.get_font_size_multiplier())
                run.font.name = self.font_name  # Ensure font consistency
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)  # For East Asian fonts

    def estimate_line_height(self, line, paragraph_data):
        # DIFF
        font_size_pt = paragraph_data.get_font_size() * self.language_config.get_font_size_multiplier()

        # Load font
        font = ImageFont.truetype(self.language_config.get_target_font_path(), size=font_size_pt)

        line_copy = str(line)
        self.remove_angle_brackets(line_copy)

        bbox = font.getbbox(line_copy)
        line_height = bbox[3] - bbox[1]
        return line_height

    @staticmethod
    def remove_angle_brackets(text: str) -> str:
        """
        Removes surrounding angle brackets like <...> but keeps inner content.
        """
        # Remove one layer of angle brackets if the entire string is wrapped in them
        cleaned = re.sub(r'^<\s*(.*?)\s*>$', r'\1', text)

        return cleaned.strip()

    @staticmethod
    def add_horizontal_line(paragraph):
        """Adds a proper horizontal line using paragraph border"""
        p_pr = paragraph._p.get_or_add_pPr()

        # Create border element if it doesn't exist
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)

        # Create bottom border with specifications - now using black color
        bottom_border = OxmlElement('w:top')
        bottom_border.set(qn('w:val'), 'single')  # Line style
        bottom_border.set(qn('w:sz'), '6')  # Line width (6ths of a point)
        bottom_border.set(qn('w:space'), '0')  # Space above line
        bottom_border.set(qn('w:color'), '000000')  # Black color

        p_bdr.append(bottom_border)

    def _add_text_with_styling(self, paragraph_obj, text, paragraph_data):
        spans = self.parse_styled_spans(text)
        for span in spans:
            run = paragraph_obj.add_run(span["text"] + " ")
            run.font.size = Pt(paragraph_data.get_font_size() * self.language_config.get_font_size_multiplier())
            run.font.name = self.font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            if span.get("bold"):
                run.bold = True

    @staticmethod
    def parse_styled_spans(text) -> List[Dict]:
        """
        Parses text and marks any content inside angle brackets <...> as bold.
        Returns a list of spans: [{"text": ..., "bold": True/False}]
        """
        spans = []
        pattern = re.compile(r'<(.*?)>')

        last_index = 0

        for match in pattern.finditer(text):
            start, end = match.span()

            # Add normal text before the tag
            if start > last_index:
                normal_text = text[last_index:start].strip()
                if normal_text:
                    spans.append({"text": normal_text, "bold": False})

            # Add bold text inside angle brackets
            bold_text = match.group(1).strip()
            if bold_text:
                spans.append({"text": bold_text, "bold": True})

            last_index = end

        # Add any remaining text after the last tag
        if last_index < len(text):
            tail = text[last_index:].strip()
            if tail:
                spans.append({"text": tail, "bold": False})

        return spans

    def _configure_docx_section(self, section):
        # Set standard page dimensions and margins
        section.page_height = Inches(Utils.STANDARDIZED_PAGE_HEIGHT)
        section.page_width = Inches(Utils.STANDARDIZED_PAGE_WIDTH)
        section.top_margin = Inches(Utils.STANDARDIZED_TOP_MARGIN)
        section.bottom_margin = Inches(Utils.STANDARDIZED_BOTTOM_MARGIN)
        section.left_margin = Inches(Utils.STANDARDIZED_LEFT_MARGIN)
        section.right_margin = Inches(Utils.STANDARDIZED_RIGHT_MARGIN)

        section.different_first_page_header_footer = True
        section.footer.is_linked_to_previous = False
        section.header.is_linked_to_previous = False

        return section

    def set_rtl(self, paragraph):
        if self.language_config.get_right_to_left():
            """Set the paragraph to Right-to-Left (RTL) direction."""
            p = paragraph._p  # Access the XML <w:p> element
            pPr = p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')  # Enable RTL
            pPr.append(bidi)

    def _set_paragraph_alignment_and_indent(self, paragraph_obj, paragraph_data, page, section,
                                            continue_para=False):
        para_format = paragraph_obj.paragraph_format
        left_indent = int(paragraph_data.get_start()) - int(page.get_min_x())
        right_indent = int(page.get_max_x()) - int(paragraph_data.get_end())
        para_start = self.document_processor.get_paragraph_start()

        if abs(left_indent - right_indent) > 0:
            if int(paragraph_data.get_start()) != int(para_start) and int(
                    paragraph_data.get_para_bbox().x0) != page.get_min_x() and int(paragraph_data.get_end()) == int(
                page.get_max_x()):
                paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para_format.first_line_indent = Inches(0)
            else:
                paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
                relative_indent_in = 0
                if not continue_para:
                    bbox_x0_in = int(paragraph_data.get_para_bbox().x0) / 72
                    relative_indent_in = (bbox_x0_in - section.left_margin.inches)
                    if int(paragraph_data.get_para_bbox().x0) == int(page.get_min_x()):
                        relative_indent_in = 0
                    elif int(paragraph_data.get_para_bbox().x0) == int(para_start):
                        relative_indent_in = relative_indent_in / 2
                para_format.first_line_indent = Inches(relative_indent_in)
        else:
            paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_format.first_line_indent = Pt(0)

    def _set_paragraph_spacing(self, paragraph_obj, paragraph_data):
        para_format = paragraph_obj.paragraph_format
        # TODO: DIFFERENT LINE SPACING FOR DIFFERENT PAGES
        para_format.line_spacing = Pt(
            paragraph_data.get_font_size() *
            self.language_config.get_font_size_multiplier() *
            self.language_config.get_line_spacing_multiplier()
        )
        para_format.space_after = Pt(paragraph_data.get_font_size() * 0.5)
        para_format.space_before = Pt(0)

    def _add_page_number(self, section, page_number):
        self.logger.info(f'Adding Page Number: {page_number}')
        """Adds footer ONLY to the current section's next page"""
        # 1. Get current section (always use last section)
        section.footer_distance = Inches(Utils.STANDARDIZED_FOOTER_DISTANCE)
        footer = section.first_page_footer
        footer.is_linked_to_previous = False

        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_format = footer_para.paragraph_format
        footer_format.space_before = Pt(2.5)
        run = footer_para.add_run(str(page_number))
        run.font.name = self.font_name  # Ensure font consistency
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)  # For East Asian fonts

    def add_page_numbers(self):
        self.logger.info('Adding Page Numbers')
        extracted_page_number, page_number_start = self.document_processor.get_page_number_info()
        if extracted_page_number is not None and page_number_start is not None:
            for idx, section in enumerate(self.document.sections):
                if idx >= page_number_start:
                    self._add_page_number(section, extracted_page_number)
                    extracted_page_number += 1
