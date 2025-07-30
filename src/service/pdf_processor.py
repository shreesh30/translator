import logging
import os
import queue
import re
from concurrent.futures import ThreadPoolExecutor
from queue import Queue
from typing import List
import threading

from PIL.ImageFont import ImageFont
from docx import Document

from src.model.language_config import LanguageConfig
from src.model.paragraph import Paragraph
from src.model.line import Line
from PIL import ImageFont
from src.model.table import Table
from src.service.document_builder import DocumentBuilder
from src.service.document_processor import DocumentProcessor
from IndicTransToolkit.processor import IndicProcessor

from src.utils.utils import Utils

logger = logging.getLogger(__name__)


class PDFProcessor:
    def __init__(self,
                 lang_config: LanguageConfig,
                 gpu_task_queue: Queue,
                 gpu_result_queue: Queue,
                 input_path: str,
                 output_path: str,
                 max_workers: int = 4):
        self.lang_config = lang_config
        self.gpu_task_queue = gpu_task_queue
        self.gpu_result_queue = gpu_result_queue
        self.input_path = input_path
        self.output_path = output_path
        self.max_workers = max_workers

    def translate_text(self, text: str) -> str:
        """Simplified translation request using a new Queue per request"""
        result_queue = Queue()  # New queue for each request

        # Put task in queue
        self.gpu_task_queue.put({
            'text': text,
            'lang': self.lang_config.target_language_key,
            'result_queue': result_queue  # Pass the queue directly
        })

        # Wait for result
        try:
            return result_queue.get(timeout=30.0)
        except queue.Empty:
            logger.error("Timeout waiting for translation result")
            return ""

    def process_file(self, filename: str):
        """Process single PDF file"""
        try:
            file_path = os.path.join(self.input_path, filename)
            logger.info(f"Processing {file_path} for {self.lang_config.target_language}")

            # 1. Extract document structure
            doc_processor = DocumentProcessor(file_path, self.lang_config)
            doc_processor.process_document()
            elements = doc_processor.get_elements()

            # 2. Parallel translation
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Submit all translation tasks
                futures = []
                for element in elements:
                    if element.type == "paragraph":
                        futures.append(executor.submit(
                            self._translate_paragraph,
                            element,
                            doc_processor
                        ))
                    elif element.type == "table":
                        futures.append(executor.submit(
                            self._translate_table,
                            element
                        ))

                # Wait for completion
                for future in futures:
                    future.result()

            # 3. Build output document
            self._build_docx(elements, filename, doc_processor)

        except Exception as e:
            logger.error(f"Failed processing {filename}: {e}")

    def _translate_paragraph(self, paragraph, doc_processor):
        """Translate paragraph content"""
        try:
            # Combine all lines for context-aware translation
            full_text = " ".join(line.get_text() for line in paragraph.get_lines())
            translated = self.translate_text(full_text)

            # Reconstruct lines with original formatting
            lines = self._split_into_lines(translated, paragraph, doc_processor)
            paragraph.set_lines(lines)

            # Process sub-paragraphs if any
            for sub_para in paragraph.get_sub_paragraphs():
                self._translate_paragraph(sub_para, doc_processor)

            return paragraph

        except Exception as e:
            logger.error(f"Paragraph translation failed: {e}")
            return paragraph

    def _translate_table(self, table):
        """Translate table content"""
        try:
            # Translate title
            if table.get_title().get_text():
                translated = self.translate_text(table.get_title().get_text())
                table.get_title().set_text(translated)

            if table.get_sub_title().get_text():
                translated = self.translate_text(table.get_sub_title().get_text())
                table.get_sub_title().set_text(translated)

            # Translate content
            for column in table.get_columns():
                for row in column:
                    if row.get_text():
                        translated = self.translate_text(row.get_text())
                        row.set_text(translated)

            return table
        except Exception as e:
            logger.error(f"Table translation failed: {e}")
            return table

    @staticmethod
    def is_tag(word: str) -> bool:
        return re.fullmatch(r'</?\w+>', word.strip()) is not None

    def _split_into_lines(self, para_text, paragraph, document_processor):
        """Splits text into lines using accurate point/inch measurements"""
        font_size = paragraph.get_font_size()

        # 1. Get font metrics in points
        font_size_pt = font_size * self.lang_config.get_font_size_multiplier()

        font = ImageFont.truetype(self.lang_config.get_target_font_path(), size=font_size_pt)

        # 2. Calculate max width IN POINTS (1 inch = 72 points)
        max_width_pt = Utils.USABLE_PAGE_WIDTH * 72

        # 3. Language-aware splitting
        return self._split_words(para_text, font, max_width_pt, paragraph, document_processor)

    def _split_words(self ,text, font, max_width_pt, paragraph, document_processor):
        """Splits space-separated languages using accurate width measurements."""
        words = text.split()
        lines = []
        current_line = []
        current_width = 0

        space_width = font.getlength(" ")
        tab_width = 4 * space_width

        para_indent = (
            tab_width if int(paragraph.get_para_bbox().x0) == document_processor.get_paragraph_start()
            else 0
        )

        applied_indent = False

        for i, word in enumerate(words):
            if self.is_tag(word):
                current_line.append(word)
                continue

            word_width = font.getlength(word)

            projected_width = current_width + space_width + word_width

            if not applied_indent:
                projected_width = word_width + para_indent
                applied_indent = True

            if projected_width > max_width_pt:
                lines.append(" ".join(current_line))
                current_line = [word]
                current_width = word_width
            else:
                current_line.append(word)
                current_width = projected_width

        if current_line:
            lines.append(" ".join(current_line))

        return lines

    def _build_docx(self, elements, filename, doc_processor):
        """Generate translated DOCX"""
        doc = Document()
        builder = DocumentBuilder(document=doc, language_config=self.lang_config, document_processor=doc_processor)
        pages = doc_processor.get_pages()
        try:
            for element in elements:
                if element.type == "paragraph":
                    builder.add_paragraph(element, pages)
                elif element.type == "table":
                    builder.add_table(element)

            output_dir = os.path.join(
                self.output_path,
                self.lang_config.target_language
            )
            os.makedirs(output_dir, exist_ok=True)

            output_path = os.path.join(output_dir, f"{os.path.splitext(filename)[0]}.docx")
            doc.save(output_path)
            logger.info(f"Saved translation to {output_path}")

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")

    def run(self):
        """Process all PDFs for this language"""
        pdf_files = [
            f for f in os.listdir(self.input_path)
            if f.lower().endswith('.pdf')
        ]

        with ThreadPoolExecutor(max_workers=2) as executor:  # 2 files at a time
            executor.map(self.process_file, pdf_files)